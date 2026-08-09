"""
bases.py — Capa C (parte 1): convierte los documentos de las bases a TEXTO, local.
NO manda imágenes a la API: extrae texto con pdfplumber y, solo si una página no tiene
capa de texto (escaneada), cae a OCR local (pytesseract). Así se paga ~1/3 de los tokens
y no se choca con el tope de 100 páginas / 32 MB del modo PDF nativo.

  documento_a_texto(ruta)  -> enruta por extensión (.pdf/.docx/.xlsx) y devuelve el mismo dict
  pdf_a_texto(ruta)        -> {"paginas":[...], "texto": str, "ocr_paginas":[...], "n_paginas":int}
  docx_a_texto(ruta)       -> idem, incluyendo TABLAS (anexos y formularios)
  xlsx_a_texto(ruta)       -> idem, hoja por hoja (Anexo económico / itemizado)
  secciones(texto)         -> dict {seccion: fragmento}  (recorte por encabezados típicos de bases)
  texto_para_analisis(ruta, max_chars) -> texto recortado a las secciones relevantes (lo que se manda a Claude)

OCR es opcional: si pytesseract/pdf2image no están, se omite y se avisa (no rompe).

Nota: hasta jul-2026 solo se leían PDFs y todo .docx/.xlsx se descartaba en silencio en
capa_c.py. Eso perdía justo los anexos económicos y los itemizados, que son los que traen
las cantidades y el presupuesto — es decir, lo que hace falta para cubicar.
"""
import os
import re
import unicodedata
from pathlib import Path
import pdfplumber


MIN_CHARS_PAGINA = 40   # bajo esto, la página se considera "sin texto" → candidata a OCR

# Dónde buscar el binario de Tesseract cuando NO está en el PATH (el instalador de Windows no lo
# agrega). Sin esto, pytesseract lanza TesseractNotFoundError, `_ocr_pagina` lo traga y devuelve ""
# → una licitación cuyo PDF de bases es un escaneo queda SIN TEXTO y es inanalizable, en silencio.
# Medido sobre un lote real: 8 de 25 licitaciones estaban en ese estado.
TESSERACT_CANDIDATOS = [
    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
    r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    "/usr/bin/tesseract", "/usr/local/bin/tesseract", "/opt/homebrew/bin/tesseract",
]


def tesseract_exe():
    """Ruta del binario de Tesseract: LOVEO_TESSERACT, o el primer candidato que exista.
    None si no hay ninguno (ahí sí el OCR no está disponible y hay que avisarlo, no tragárselo)."""
    forzado = os.environ.get("LOVEO_TESSERACT", "").strip()
    if forzado:
        return forzado if Path(forzado).exists() else None
    return next((p for p in TESSERACT_CANDIDATOS if Path(p).exists()), None)


def ocr_disponible():
    """True si se puede OCRear de verdad (pytesseract importable + binario encontrado o en PATH)."""
    try:
        import pytesseract
    except ImportError:
        return False
    exe = tesseract_exe()
    if exe:
        pytesseract.pytesseract.tesseract_cmd = exe
        return True
    try:                                    # sin candidato: puede estar en el PATH
        pytesseract.get_tesseract_version()
        return True
    except Exception:
        return False

# Encabezados típicos en bases de Mercado Público → secciones que alimentan el score y el Q&A.
SECCIONES = {
    "presupuesto":   r"(presupuesto|monto\s+disponible|monto\s+estimado|valor\s+referencial|financiamiento)",
    "plazo":         r"(plazo\s+de\s+entrega|plazo\s+de\s+ejecuci|plazo\s+del\s+contrato|plazo\s+de\s+fabricaci)",
    "garantias":     r"(garant[ií]a|boleta|seriedad\s+de\s+la\s+oferta|fiel\s+cumplimiento)",
    "evaluacion":    r"(criterios?\s+de\s+evaluaci|pauta\s+de\s+evaluaci|ponderaci)",
    "requisitos":    r"(requisitos|antecedentes\s+t[eé]cnicos|experiencia|profesional\s+residente|admisibilidad)",
    "visita":        r"(visita\s+a\s+terreno|visita\s+t[eé]cnica)",
    "subcontrato":   r"(subcontrat)",
    "anticipo":      r"(anticipo|estado\s+de\s+pago|forma\s+de\s+pago)",
}




def _ocr_pagina(page):
    """OCR opcional de una página pdfplumber. Devuelve texto o '' si no hay OCR disponible.
    Apunta pytesseract al binario encontrado por `tesseract_exe()` antes de llamarlo: en Windows
    el instalador no toca el PATH y sin esto el OCR fallaba en silencio (ver TESSERACT_CANDIDATOS)."""
    try:
        import pytesseract
        exe = tesseract_exe()
        if exe:
            pytesseract.pytesseract.tesseract_cmd = exe
        im = page.to_image(resolution=200).original  # PIL Image
        return pytesseract.image_to_string(im, lang="spa")
    except Exception:
        return ""


def pdf_a_texto(ruta):
    ruta = str(ruta)
    paginas, ocr_pgs = [], []
    with pdfplumber.open(ruta) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            t = page.extract_text() or ""
            if len(t.strip()) < MIN_CHARS_PAGINA:        # página sin capa de texto → OCR
                ocr = _ocr_pagina(page)
                if ocr.strip():
                    t = ocr
                    ocr_pgs.append(i)
            paginas.append(t)
    return {"paginas": paginas, "texto": "\n".join(paginas),
            "ocr_paginas": ocr_pgs, "n_paginas": len(paginas)}


def docx_a_texto(ruta):
    """Texto de un .docx, incluidas las TABLAS (ahí viven los anexos y los itemizados)."""
    from docx import Document
    d = Document(str(ruta))
    out = [p.text for p in d.paragraphs if p.text.strip()]
    for t in d.tables:
        for fila in t.rows:
            celdas = [c.text.strip() for c in fila.cells]
            if any(celdas):
                out.append(" | ".join(celdas))
    return {"paginas": out, "texto": "\n".join(out), "ocr_paginas": [], "n_paginas": len(out)}


def xlsx_a_texto(ruta):
    """Texto de un .xlsx/.xlsm — el Anexo económico y el itemizado casi siempre vienen así."""
    from openpyxl import load_workbook
    wb = load_workbook(str(ruta), data_only=True)
    out = []
    for ws in wb.worksheets:
        out.append(f"=== HOJA: {ws.title} ===")
        for fila in ws.iter_rows(values_only=True):
            celdas = ["" if v is None else str(v).strip() for v in fila]
            if any(celdas):
                out.append(" | ".join(celdas))
    return {"paginas": out, "texto": "\n".join(out), "ocr_paginas": [], "n_paginas": len(wb.worksheets)}


# Comprimidos: los compradores suben la planimetría y los anexos dentro de .zip/.rar. Sin esto
# quedaban fuera del análisis expedientes completos (p.ej. los 3 anexos de la U. de Atacama).
COMPRIMIDOS = (".zip", ".rar", ".7z")
_SIETE_ZIP = [r"C:\Program Files\7-Zip\7z.exe", r"C:\Program Files (x86)\7-Zip\7z.exe", "7z"]
MAX_ARCHIVOS_COMPRIMIDO = 40      # cota defensiva ante un comprimido con miles de entradas


def _bin_7z():
    import shutil
    for cand in _SIETE_ZIP:
        if Path(cand).is_file() or shutil.which(cand):
            return cand
    return None


def comprimido_a_texto(ruta):
    """Extrae un .zip/.rar/.7z a un directorio temporal y lee los documentos soportados de adentro.

    Requiere 7-Zip instalado (gratis). Si no está, lanza para que capa_c lo registre en `meta`
    en vez de descartar el archivo en silencio."""
    import tempfile, subprocess
    exe = _bin_7z()
    if not exe:
        raise RuntimeError("7-Zip no está instalado: no se puede abrir el comprimido "
                           f"{Path(ruta).name}. Instalar de 7-zip.org (gratis).")
    partes, leidos = [], 0
    with tempfile.TemporaryDirectory() as tmp:
        r = subprocess.run([exe, "x", "-y", f"-o{tmp}", str(ruta)],
                           capture_output=True, text=True, errors="replace", timeout=300)
        if r.returncode != 0:
            raise RuntimeError(f"7-Zip no pudo extraer {Path(ruta).name}: {(r.stderr or '')[:120]}")
        for p in sorted(Path(tmp).rglob("*")):
            if leidos >= MAX_ARCHIVOS_COMPRIMIDO:
                partes.append(f"[... comprimido truncado en {MAX_ARCHIVOS_COMPRIMIDO} archivos]")
                break
            ext = p.suffix.lower()
            if not p.is_file() or ext not in LECTORES_SIMPLES:
                continue
            try:
                partes.append(f"### {p.name}\n" + LECTORES_SIMPLES[ext](p)["texto"])
                leidos += 1
            except Exception as e:
                partes.append(f"### {p.name}\n[no legible: {type(e).__name__}]")
    txt = "\n\n".join(partes)
    return {"paginas": partes, "texto": txt, "ocr_paginas": [], "n_paginas": leidos}


def texto_plano_a_texto(ruta):
    """.txt/.csv/.md: se leen tal cual. Aparecen como anexos sueltos (listados de items,
    aclaraciones pegadas del foro) y antes caían al fallback de pdfplumber, que fallaba."""
    raw = Path(str(ruta)).read_bytes()
    for enc in ("utf-8", "cp1252", "latin-1"):        # los pliegos chilenos vienen en los tres
        try:
            txt = raw.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        txt = raw.decode("utf-8", "replace")
    return {"paginas": [txt], "texto": txt, "ocr_paginas": [], "n_paginas": 1}


def _legacy_office(ruta):
    """.doc/.xls del Office viejo: no hay lector puro-Python confiable. Se avisa qué falta
    en vez de descartar el archivo en silencio."""
    raise RuntimeError(
        f"{Path(ruta).suffix} (Office antiguo) no es legible sin conversor. Instalar LibreOffice "
        f"y convertir con: soffice --headless --convert-to docx \"{Path(ruta).name}\"")


# Lectores de archivo suelto (los que se pueden usar DENTRO de un comprimido, sin recursión).
LECTORES_SIMPLES = {".pdf": pdf_a_texto, ".docx": docx_a_texto,
                    ".xlsx": xlsx_a_texto, ".xlsm": xlsx_a_texto,
                    ".txt": texto_plano_a_texto, ".csv": texto_plano_a_texto,
                    ".md": texto_plano_a_texto}

# Extensión → lector. Lo que no esté acá sigue cayendo a pdfplumber (y falla explícito).
LECTORES = dict(LECTORES_SIMPLES,
                **{e: comprimido_a_texto for e in COMPRIMIDOS},
                **{".doc": _legacy_office, ".xls": _legacy_office})


# Un OCR fallido no devuelve vacío: devuelve BASURA ("HoRADENCIO | 0% — —", "Ol uates GonzZa lez").
# Por eso no sirve medir cantidad de texto sino CALIDAD: qué fracción de los tokens parece una
# palabra de verdad. Medido sobre expedientes reales: acta manuscrita 0,29 · documentos legibles
# 0,55-0,58. El corte en 0,45 los separa con holgura.
CALIDAD_MINIMA = 0.45
MIN_TOKENS_PARA_JUZGAR = 60          # con menos tokens la proporción es ruido
RX_DOC_CRITICO = re.compile(r"acta|visita|apertura|adjudicaci", re.I)
_RX_PALABRA = re.compile(r"[A-Za-zÁÉÍÓÚÜÑáéíóúüñ]{3,20}")


_RX_CASE_RARA = re.compile(r"[a-záéíóúüñ][A-ZÁÉÍÓÚÜÑ]")   # 'GonzZa', 'HoRADENCIO': mayúscula intercalada


def _parece_palabra(w):
    """Palabra plausible: solo letras, largo razonable y sin mayúscula intercalada.
    La mayúscula en medio de la palabra es la huella más clara del OCR sobre manuscrito."""
    return bool(_RX_PALABRA.fullmatch(w)) and not _RX_CASE_RARA.search(w)


def calidad_texto(texto):
    """De los tokens que CONTIENEN LETRAS, qué fracción es una palabra plausible.

    Se ignoran números, montos y separadores de tabla a propósito: un itemizado legítimo está
    lleno de cifras y '|', y penalizarlo lo marcaría como ilegible. Lo que delata al OCR fallido
    son los tokens con letras que no son palabras ('HoRADENCIO', 'GonzZa', 'I$pñ]')."""
    con_letras = [w for w in (texto or "").split() if any(c.isalpha() for c in w)]
    if len(con_letras) < MIN_TOKENS_PARA_JUZGAR:
        return 1.0                    # muy corto para juzgar: no se marca
    return sum(1 for w in con_letras if _parece_palabra(w)) / len(con_letras)


def necesita_vision(ruta, res):
    """True si el texto extraído es ilegible aunque el OCR haya corrido.

    El caso real: el acta de visita a terreno de una licitación con visita EXCLUYENTE define
    quién puede ofertar y suele estar manuscrita — Tesseract devuelve basura, no vacío. Es el
    escenario donde conviene gastar visión; `es_critico()` marca cuándo además urge."""
    return calidad_texto(res.get("texto")) < CALIDAD_MINIMA


def es_critico(ruta):
    """El nombre del archivo sugiere que decide la admisibilidad (acta de visita/apertura)."""
    return bool(RX_DOC_CRITICO.search(Path(str(ruta)).name))


def _firma(ruta):
    """Identidad barata del archivo: tamaño + mtime. Si cambia, el caché se invalida."""
    st = Path(str(ruta)).stat()
    return f"{st.st_size}:{int(st.st_mtime)}"


def _cache_leer(ruta):
    """Texto ya extraído para este archivo, si la firma coincide. Silencioso ante cualquier
    problema de BD: el caché es una optimización, nunca una dependencia dura."""
    try:
        import db, json as _json
        with db.conn() as c:
            row = c.execute("SELECT texto_cache, texto_firma, texto_ocr_paginas FROM documentos "
                            "WHERE ruta_local=? AND texto_cache IS NOT NULL", (str(ruta),)).fetchone()
        if row and row[1] == _firma(ruta):
            return {"paginas": [], "texto": row[0],
                    "ocr_paginas": _json.loads(row[2] or "[]"), "n_paginas": 0}
    except Exception:
        pass
    return None


def _cache_guardar(ruta, res):
    try:
        import db, json as _json, datetime
        with db.conn() as c:
            c.execute("UPDATE documentos SET texto_cache=?, texto_firma=?, texto_ocr_paginas=?, "
                      "texto_at=? WHERE ruta_local=?",
                      (res["texto"], _firma(ruta), _json.dumps(res["ocr_paginas"]),
                       datetime.datetime.now().isoformat(timespec="seconds"), str(ruta)))
    except Exception:
        pass


def documento_a_texto(ruta, usar_cache=True):
    """Punto de entrada único: enruta por extensión. Los anexos y itemizados vienen en
    .docx/.xlsx y antes se descartaban en silencio, perdiendo cantidades y presupuesto.

    Cachea el resultado en `documentos.texto_cache` (clave: tamaño+mtime del archivo). El OCR
    de un pliego escaneado son minutos; sin caché se repetía en cada corrida de Capa C."""
    if usar_cache:
        hit = _cache_leer(ruta)
        if hit is not None:
            return hit
    ext = Path(str(ruta)).suffix.lower()
    res = LECTORES.get(ext, pdf_a_texto)(ruta)
    if usar_cache:
        _cache_guardar(ruta, res)
    return res


def _normalizar_con_mapa(s):
    """(texto_normalizado, mapa) para buscar encabezados y poder volver al texto ORIGINAL.

    `textnorm.na` (encode ascii+ignore) BORRA lo no representable (viñetas, dashes, comillas
    tipográficas, °…) y las ligaduras ('ﬁ') expanden a 2 chars: en ambos casos los offsets del
    normalizado NO coinciden con los del original, y recortar con ellos devolvía secciones corridas.
    Acá se normaliza sin restricción de longitud y `mapa[i]` guarda el índice ORIGINAL que produjo
    el char i del normalizado, así el recorte siempre cae en el lugar correcto."""
    out, mapa = [], []
    for i, ch in enumerate(s or ""):
        d = unicodedata.normalize("NFKD", ch)
        base = "".join(c for c in d if not unicodedata.combining(c)) or ch
        for c in base.lower():                    # 'ﬁ' → 'f','i' (ambos apuntan al mismo índice)
            out.append(c)
            mapa.append(i)
    mapa.append(len(s or ""))                     # centinela: fin del texto
    return "".join(out), mapa


# Texto que se repite IGUAL en toda licitación de Mercado Público y no aporta nada al análisis:
# recitales de la Ley 19.886, listado de inhabilidades, pies de firma electrónica, numeración de
# páginas. Medido sobre expedientes reales: ~4% del texto, y sube en los pliegos largos.
BOILERPLATE = re.compile(
    r"firmado electr[oó]nicamente de acuerdo con la ley|"
    r"doc\.digital\.gob\.cl/validador|"
    r"para verificar la integridad y autenticidad de este documento|"
    r"p[aá]gina\s+\d+\s+de\s+\d+|"
    r"^\s*www\.[\w.\-]+\s*$|"
    r"haber sido condenado por cualquiera de los delitos de cohecho|"
    r"registrar (una o m[aá]s deudas tributarias|deudas previsionales)|"
    r"la presentaci[oó]n al registro nacional de proveedores de uno o m[aá]s documentos falsos|"
    r"haber sido declarado en quiebra por resoluci[oó]n judicial|"
    r"haber sido condenado por pr[aá]cticas antisindicales",
    re.I | re.M)


def sin_boilerplate(texto):
    """Saca las líneas de relleno legal idénticas en toda licitación. No usa API."""
    return "\n".join(l for l in (texto or "").split("\n") if not BOILERPLATE.search(l))


def dedup_lineas(bloques, min_len=40):
    """Quita las líneas repetidas ENTRE documentos de una misma licitación: el decreto repite
    las bases, los anexos repiten los formularios. Medido: ~11% del expediente. La primera
    aparición se conserva; las siguientes se descartan. No usa API."""
    vistas, out = set(), []
    for b in bloques:
        keep = []
        for l in (b or "").split("\n"):
            k = " ".join(l.split()).lower()
            if len(k) >= min_len:
                if k in vistas:
                    continue
                vistas.add(k)
            keep.append(l)
        out.append("\n".join(keep))
    return out


# Marcas de que un tramo trae la sección DE VERDAD y no la línea del índice:
# porcentajes, montos, puntajes, tramos. El índice de un pliego no tiene nada de esto.
_SENAL = re.compile(r"\d{1,3}\s?%|\$\s?[\d\.]{6,}|puntaje|puntos|ponderaci|"
                    r"d[ií]as\s+(corridos|h[aá]biles)|utm|inadmisib", re.I)

VENTANA_SENAL = 1800    # chars posteriores al match donde se mide la densidad de señal
_RX_INDICE = re.compile(r"\.{4,}")     # "Criterios de evaluacion ........ pagina 12"


def _es_linea_de_indice(n, pos):
    """La línea del match tiene puntos guía: es una entrada del índice, no la sección."""
    ini = n.rfind("\n", 0, pos) + 1
    fin = n.find("\n", pos)
    return bool(_RX_INDICE.search(n[ini:fin if fin != -1 else len(n)]))


def _densidad_senal(n, pos):
    """Cuántas marcas de decisión hay justo después de `pos`. Sirve para distinguir el
    encabezado real de su mención en el índice o en una referencia cruzada."""
    return len(_SENAL.findall(n[pos:pos + VENTANA_SENAL]))


def secciones(texto):
    """Parte el texto por encabezados relevantes; cada sección = desde su match hasta el siguiente.
    Busca sobre el texto normalizado y traduce las posiciones al ORIGINAL vía el mapa de offsets.

    De cada encabezado se toma la ocurrencia con MÁS señal alrededor, no la primera: en los
    pliegos con índice, 'criterios de evaluación' aparece primero en el índice y antes se
    capturaba esa línea en vez de la matriz de evaluación real."""
    n, mapa = _normalizar_con_mapa(texto)
    hits = []
    for nombre, pat in SECCIONES.items():
        mejor, mejor_score = None, -1
        for m in re.finditer(pat, n):
            if _es_linea_de_indice(n, m.start()):     # entrada del índice: nunca es la sección
                continue
            s = _densidad_senal(n, m.start())
            if s >= mejor_score:                      # empate → gana la ÚLTIMA: el índice va primero
                mejor, mejor_score = m.start(), s
        if mejor is None:                             # todas eran índice → usar la primera igual
            m0 = re.search(pat, n)
            mejor = m0.start() if m0 else None
        if mejor is not None:
            hits.append((mejor, nombre))
    hits.sort()
    out = {}
    for idx, (pos, nombre) in enumerate(hits):
        ini = mapa[pos]                                     # posición en el texto ORIGINAL
        fin = (mapa[hits[idx + 1][0]] if idx + 1 < len(hits)
               else min(len(texto), ini + 2500))
        out[nombre] = texto[ini:fin].strip()
    return out


# --------------------------------------------------------------------------------------------
# Visita a terreno detectada desde el TEXTO de las bases.
#
# Por qué existe esto: el API de Mercado Público devuelve `Fechas.FechaVisitaTerreno` en null
# prácticamente siempre (verificado sobre 10 licitaciones con visita declarada en sus bases), así
# que la alerta de alertas.py nunca se disparaba. La visita sí está escrita en el pliego, y cuando
# es excluyente define quién puede ofertar: no detectarla costó una licitación de $259 M.
# --------------------------------------------------------------------------------------------
# "visita a terreno", "visita de terreno", "visita técnica" y las variantes pegadas que deja el
# OCR o la extracción de tablas ("VisitadeTerreno").
RX_VISITA = re.compile(r"visita\s*(?:a|de)?\s*(?:terreno|t[eé]cnica)", re.I)
RX_PALABRA_VISITA = re.compile(r"visitas?\b", re.I)   # incluye "esta visita", "dicha visita"
RX_OBLIGATORIA = re.compile(
    r"visita[^.]{0,80}(obligatori|excluyente)|"
    r"(obligatori|excluyente)[^.]{0,80}visita|"
    r"no\s+(concurr|asist|particip)[^.]{0,120}(inadmisib|excluid|fuera\s+de\s+bases|rechaz)|"
    r"(inadmisib|excluid)[^.]{0,120}no\s+(concurr|asist|particip)", re.I)
# Los pliegos escriben el puntaje en los dos órdenes: "…a la visita 100 puntos" y
# "…no participe tendrá puntaje 0". Hay que cubrir ambos.
RX_PUNTUA = re.compile(r"(asist|particip|concurr)[^.]{0,80}(\d{1,3}\s*puntos?|puntajes?\s*\d{1,3})|"
                       r"(\d{1,3}\s*puntos?|puntajes?\s*\d{1,3})[^.]{0,80}(asist|particip|concurr)",
                       re.I)
# OJO con los \b: sin ellos, "\bno\b" llega a matchear el final de "terreNO Se realizará" y una
# visita obligatoria se leia como prohibida. Todos los "no" van anclados a palabra completa.
_NEGACION = (r"(?:prohibid\w*|\bno\s+se\s+contempla\w*|\bno\s+contemplan?\b|\bno\s+habr[aá]\b|"
             r"\bno\s+se\s+(?:realizar|efectuar)\w*)")
RX_PROHIBIDA = re.compile(rf"visitas?[^.]{{0,80}}{_NEGACION}|{_NEGACION}[^.]{{0,80}}visitas?", re.I)
# "Cuando se disponga segun las Bases visita a terreno con caracter obligatorio…" es una clausula
# CONDICIONAL de inadmisibilidad, boilerplate en muchos pliegos: no significa que ESTA licitacion
# tenga visita. Sin este filtro, Cabildo (que dice expresamente "No se contempla") daba EXCLUYENTE.
RX_CONDICIONAL = re.compile(r"(cuando\s+se\s+disponga|en\s+caso\s+(de\s+)?que\s+se\s+disponga|"
                            r"si\s+(se\s+)?(hubiere|hubiese|existiere)|de\s+haberla)", re.I)
# "3 días corridos posteriores a la publicación" | "el martes 4 de agosto de 2026" | "10:00 horas"
RX_FECHA_REL = re.compile(r"(?:al\s+)?(?:el\s+)?(\w+|\d{1,2})\s*\(?\d*\)?\s*d[ií]as?\s+"
                          r"(?:corridos?|h[aá]biles)[^.]{0,60}publicaci[oó]n", re.I)
RX_FECHA_ABS = re.compile(r"\d{1,2}\s+de\s+[a-záéíóú]+\s+(?:de\s+)?\d{4}|"
                          r"\d{1,2}[/-]\d{1,2}[/-]\d{2,4}", re.I)
RX_HORA = re.compile(r"\d{1,2}[:.]\d{2}\s*(?:hrs?|horas)?", re.I)


def visita_a_terreno(texto):
    """Detecta si el pliego contempla visita a terreno y con qué carácter.

    Devuelve {"menciona", "obligatoria", "puntuada", "prohibida", "fecha_literal", "fragmento"}.
    No usa API: es regex sobre el texto ya extraído.

    Fiabilidad (validado contra 7 expedientes reales de agosto 2026): la CLASIFICACIÓN
    —obligatoria / puntuada / prohibida / voluntaria / sin visita— acertó 7 de 7. La
    `fecha_literal` es ORIENTATIVA: en pliegos con muchas fechas puede tomar la de la resolución
    en vez de la de la visita. Por eso se devuelve también `fragmento`, para leer el texto real."""
    t = texto or ""
    m = RX_VISITA.search(t)
    if not m:
        return {"menciona": False, "obligatoria": False, "puntuada": False,
                "prohibida": False, "fecha_literal": None, "fragmento": ""}
    # Confirmado que el pliego tiene visita a terreno, el CARÁCTER se busca en ventanas alrededor
    # de CADA mención de la palabra "visita" (no solo de "visita a terreno"): el artículo que la
    # hace excluyente suele decir "esta visita". Anclar a la palabra evita el falso positivo de
    # las cláusulas de anexos ("...inadmisible y no participará de la evaluación"), que no hablan
    # de la visita pero calzan con el patrón si se busca sobre todo el texto.
    todo = "\n".join(" ".join(t[max(0, c.start() - 250): c.start() + 500].split())
                     for c in RX_PALABRA_VISITA.finditer(t))
    # La FECHA, en cambio, se toma de la mención más informativa (la que trae fechas cerca).
    mejor, mejor_peso = m, -1
    for cand in RX_VISITA.finditer(t):
        v = " ".join(t[max(0, cand.start() - 250): cand.start() + 700].split())
        peso = len(RX_FECHA_ABS.findall(v)) * 2 + len(RX_FECHA_REL.findall(v))
        if peso > mejor_peso:
            mejor, mejor_peso = cand, peso
    frag = " ".join(t[max(0, mejor.start() - 250): mejor.start() + 700].split())
    fecha = RX_FECHA_ABS.search(frag) or RX_FECHA_REL.search(frag)
    hora = RX_HORA.search(frag)
    literal = None
    if fecha:
        literal = fecha.group(0) + (f" · {hora.group(0)}" if hora else "")
    prohibida = bool(RX_PROHIBIDA.search(todo))
    # Una mencion de "obligatoria" solo cuenta si NO viene de una clausula condicional
    obligatoria = any(not RX_CONDICIONAL.search(
                          todo[max(0, mm.start() - 120): mm.start() + 40])
                      for mm in RX_OBLIGATORIA.finditer(todo))
    # Si el pliego dice expresamente que no contempla visita, eso manda sobre el boilerplate.
    if prohibida:
        obligatoria = False
    return {"menciona": True,
            "obligatoria": obligatoria,
            "puntuada": bool(RX_PUNTUA.search(todo)) and not prohibida,
            "prohibida": prohibida,
            "fecha_literal": literal,
            "fragmento": frag[:400]}


def texto_para_analisis(ruta, max_chars=12000):
    """Texto recortado a las secciones relevantes — esto es lo que se manda a Claude (no el PDF)."""
    full = documento_a_texto(ruta)
    secs = secciones(full["texto"])
    if secs:
        armado = "\n\n".join(f"### {k.upper()}\n{v}" for k, v in secs.items())
    else:
        armado = full["texto"]          # sin encabezados reconocidos → texto completo (se recorta abajo)
    return {"texto": armado[:max_chars], "n_paginas": full["n_paginas"],
            "ocr_paginas": full["ocr_paginas"], "secciones_halladas": list(secs.keys())}


if __name__ == "__main__":
    import sys, json
    ruta = sys.argv[1] if len(sys.argv) > 1 else "calibration/ficha_real_2450-56-LE26.pdf"
    r = pdf_a_texto(ruta)
    print(f"Páginas: {r['n_paginas']} | OCR usado en: {r['ocr_paginas'] or 'ninguna'} | chars: {len(r['texto'])}")
    secs = secciones(r["texto"])
    print("Secciones detectadas:", list(secs.keys()))
    print("\n--- muestra (primeros 600 chars) ---\n", r["texto"][:600])
